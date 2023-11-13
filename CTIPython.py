import streamlit as st
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO

def replace_placeholders(docx_template, replacements):
    for key, value in replacements.items():
        value = str(value)
        for paragraph in docx_template.paragraphs:
            for run in paragraph.runs:
                run.text = run.text.replace(f"{{{{{key}}}}}", value)

def convert_docx_to_pdf(docx_template_path):
    docx_file = Document(docx_template_path)

    pdf_content = BytesIO()
    can = canvas.Canvas(pdf_content, pagesize=letter)

    for paragraph in docx_file.paragraphs:
        style = paragraph.style
        if style and style.name:
            style_name = str(style.name)
            if style_name == "Normal":
                font_size = 12
            else:
                font_size = 14
            can.setFont("Helvetica", font_size)
            can.setFillColorRGB(0, 0, 0)  # black text color
            for run in paragraph.runs:
                can.drawString(15, 800, run.text)  # Adjust the coordinates as needed
            can.showPage()

    can.save()
    pdf_content.seek(0)

    return pdf_content.read()

# Streamlit app
st.title("Generate PDF Report")

# Input fields
docx_template_path = st.selectbox("Select a DOCX Template", ["CTAdvisory-TLP-CLEAR-Template.docx", "CTVulnerability-TLP-CLEAR-Template.docx"])
docx_template = Document(docx_template_path)
title = st.text_input("Title")
tlp_options = ["CLEAR", "RED", "AMBER+STRICT", "AMBER", "GREEN"]
tlp_value = st.selectbox("Select TLP", tlp_options)
summary = st.text_area("Summary")
notedata = st.text_area("Note")
content = st.text_area("Content")
technical_details = st.text_area("Technical Details")
iocs = st.text_area("Indicators of Compromise")
detection = st.text_area("Detection")
mitre_attack_reference = st.text_area("MITRE ATT&CK References")
mitigation = st.text_area("Mitigation")
valid_security_controls = st.text_area("Valid Security Controls")
resources = st.text_area("Resources")
reporting = st.text_area("Reporting")
disclaimer = st.text_area("Disclaimer")
references = st.text_area("References")

if st.button("Generate PDF"):
    replacements = {
        "title": title,
        "tlp_value": tlp_value,
        "summary": summary,
        "notedata": notedata,
        "content": content,
        "technical_details": technical_details,
        "iocs": iocs,
        "detection": detection,
        "mitre_attack_reference": mitre_attack_reference,
        "mitigation": mitigation,
        "valid_security_controls": valid_security_controls,
        "resources": resources,
        "reporting": reporting,
        "disclaimer": disclaimer,
        "references": references,
    }

    # Replace placeholders in the DOCX template
    replace_placeholders(docx_template, replacements)

    # Save the modified DOCX
    output_path_docx = f"Generated-Report-{title}.docx"
    docx_template.save(output_path_docx)

    # Convert DOCX to PDF
    pdf_content = convert_docx_to_pdf(output_path_docx)
    output_path_pdf = f"Generated-Report-{title}.pdf"

    with open(output_path_pdf, "wb") as pdf_file:
        pdf_file.write(pdf_content)

    st.success("PDF report generated successfully!")
    st.markdown(f"Download your report [here](/{output_path_pdf})")
